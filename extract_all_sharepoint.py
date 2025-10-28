# -*- coding: utf-8 -*-
"""
Extract All SharePoint Content and Create Word Document
"""

from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os

load_dotenv()

print("=" * 60)
print("EXTRACT ALL SHAREPOINT CONTENT")
print("=" * 60)

try:
    from app.sharepoint_selenium_extractor import extract_sharepoint_pages
    
    print("[*] Starting full SharePoint extraction...")
    print("[*] This will take a few minutes...\n")
    
    # Extract all pages
    documents = extract_sharepoint_pages()
    
    print("\n" + "=" * 60)
    print("EXTRACTION RESULTS")
    print("=" * 60)
    print(f"✅ Extracted {len(documents)} documents\n")
    
    # Create Word document with all content
    print("[*] Creating Word document with all extracted content...")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('SharePoint Content - Complete Extraction', 0)
    subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    subtitle.alignment = 1
    
    doc.add_paragraph()
    
    # Summary
    summary = doc.add_paragraph(f'Total pages extracted: {len(documents)}')
    summary.alignment = 1
    
    doc.add_page_break()
    
    # Add each document's content
    for i, document in enumerate(documents, 1):
        print(f"   Adding document {i}/{len(documents)}: {document.metadata.get('page_title')}")
        
        # Add heading with page title
        doc.add_heading(f'{i}. {document.metadata.get("page_title", "Untitled")}', 1)
        
        # Add URL
        url_para = doc.add_paragraph(f'URL: {document.metadata.get("page_url", "N/A")}')
        url_para.style = 'Intense Quote'
        
        # Add content
        content = document.page_content
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
        
        # Add page break if not last
        if i < len(documents):
            doc.add_page_break()
    
    # Save document
    output_file = 'SharePoint_Complete_Content.docx'
    doc.save(output_file)
    
    print(f"\n✅ Word document created: {output_file}")
    print(f"   Location: {os.path.abspath(output_file)}")
    print(f"   Total pages in document: {len(documents)}")
    
    # Show summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"Total documents extracted: {len(documents)}")
    print(f"\nDocuments:")
    for i, doc_item in enumerate(documents, 1):
        print(f"{i}. {doc_item.metadata.get('page_title', 'Untitled')} ({len(doc_item.page_content)} chars)")
    
    print("\n" + "=" * 60)
    print("NEXT STEPS")
    print("=" * 60)
    print("1. Review the Word document: SharePoint_Complete_Content.docx")
    print("2. If content looks good, add to vectorstore:")
    print("   - Set ENABLE_SHAREPOINT_SOURCE=true in .env")
    print("   - Set INITIALIZE_VECTORSTORE=true in .env")
    print("   - Run: python server.py")
    print("3. Your chatbot will have all SharePoint knowledge!")

except ImportError as e:
    print(f"[ERROR] Import failed: {e}")
    print("[INFO] Make sure Selenium is installed:")
    print("   pip install selenium webdriver-manager")
except Exception as e:
    print(f"[ERROR] Extraction failed: {e}")
    import traceback
    traceback.print_exc()

