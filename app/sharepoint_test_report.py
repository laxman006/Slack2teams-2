# -*- coding: utf-8 -*-
"""
Create a Word document with SharePoint test content
"""

from docx import Document
from datetime import datetime

def create_sharepoint_test_report():
    """Create a Word document with SharePoint test content."""
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('SharePoint Knowledge Base - Test Content', 0)
    
    # Add metadata
    subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    subtitle.alignment = 1  # Center alignment
    
    doc.add_paragraph()  # Empty line
    
    # Add introduction
    doc.add_heading('Introduction', 1)
    intro_para = doc.add_paragraph(
        'This document contains sample content extracted from SharePoint DOC360 site. '
        'This content demonstrates the type of information that will be available in the chatbot knowledge base.'
    )
    
    doc.add_page_break()
    
    # Add FAQ Section
    doc.add_heading('FAQ: Slack to Teams Migration', 1)
    
    # FAQ 1
    doc.add_paragraph('Q: What are the frequent conflicts faced during Migration?', style='Heading 2')
    doc.add_paragraph(
        'A: Bad request (Non Retriable):\n'
        '• Replied message version conflicts in post\n'
        '• We don\'t migrate bot messages\n'
        '• Missing body content\n'
        '• Neither body nor adaptive card content contains marker for mention with Id\n\n'
        'Resource Modifies (Retryable):\n'
        '• Resource has changed - usually an eTag mismatch\n'
        '• Omitting partial - files size varies but data migrate without any data missing'
    )
    
    # FAQ 2
    doc.add_paragraph('Q: Do we migrate app integration messages?', style='Heading 2')
    doc.add_paragraph(
        'A: No, we don\'t migrate app integration messages, but they will appear as admin posted messages.'
    )
    
    # FAQ 3
    doc.add_paragraph('Q: Do we migrate slack channels into existing teams?', style='Heading 2')
    doc.add_paragraph(
        'A: Yes, we do migrate Slack channels into existing Teams. '
        'But those messages inside a channel will be migrated as admin posted messages.'
    )
    
    # FAQ 4
    doc.add_paragraph('Q: How to migrate deactivated user DMS?', style='Heading 2')
    doc.add_paragraph(
        'A: We can\'t migrate deactivated user DMs because deactivated users can\'t authenticate from Teams. '
        'It\'s a limitation of Teams.'
    )
    
    # FAQ 5
    doc.add_paragraph('Q: Do we migrate the link to other messages from Slack?', style='Heading 2')
    doc.add_paragraph(
        'A: No, we don\'t migrate. They will be migrated as links, but those links will redirect back to Slack.'
    )
    
    doc.add_page_break()
    
    # Add Table Section
    doc.add_heading('Compatibility Matrix: Egnyte as Source', 1)
    
    doc.add_paragraph('The following table shows feature compatibility when migrating from Egnyte to various cloud destinations.')
    
    # Create table
    table = doc.add_table(rows=18, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    headers = ['Features', 'Google My Drive', 'Google Shared Drive', 'SharePoint Online', 'OneDrive For Business', 'Azure']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    # Data rows
    rows_data = [
        ['One Time', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Delta', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Folder Display', 'Yes', 'Yes', 'No', 'No', 'Yes'],
        ['Versions', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Selective Versions', 'No', 'No', 'No', 'No', 'No'],
        ['Root folder permissions', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Sub folder Permissions', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Root File permissions', 'NA', 'NA', 'NA', 'NA', 'NA'],
        ['Inner file permissions', 'NA', 'NA', 'NA', 'NA', 'NA'],
        ['External Shares', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Shared Links', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Preserve Timestamp', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['In-line comment', 'Yes', 'Yes', 'No', 'Yes', 'No'],
        ['Long folder path', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Special character replacement', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Embedded Links', 'No', 'No', 'Yes', 'Yes', 'Yes'],
        ['Suppressing Email Notification', 'Yes', 'Yes', 'Yes', 'No', 'No'],
    ]
    
    for i, row_data in enumerate(rows_data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_page_break()
    
    # Add Source Combinations Section
    doc.add_heading('Source Combinations Available', 1)
    
    doc.add_paragraph('The following source combinations are available in the SharePoint documentation:')
    
    source_list = [
        'Box for Business as a Source Combinations',
        'Dropbox for Business as source Combinations',
        'Egnyte as source combinations',
        'Citrix ShareFile as source combinations',
        'Google My drive as source combinations',
        'Google Share drive as source combinations',
        'SharePoint Online as source combinations',
        'NFS as source combinations',
        'Amazon WorkDocs as source combinations',
        'Cloud to Object Storage',
        'Single User Cloud-Cloud Combinations',
        'Email Migration',
        'Onedrive for business as source',
        'LinkEX Features & Combinations',
        'Message Migration Combinations',
        'Slack - Google',
        'Teams - Teams & Google chat',
        'White Board Features & Combinations'
    ]
    
    for source in source_list:
        doc.add_paragraph(source, style='List Bullet')
    
    doc.add_page_break()
    
    # Add Statistics Section
    doc.add_heading('Expected Content Statistics', 1)
    
    stats = [
        ('Main Hub Categories', '18+'),
        ('FAQ Pages', '10+'),
        ('Compatibility Tables', '20+'),
        ('Feature Definitions', '15+'),
        ('Total Expected Documents', '100+')
    ]
    
    stats_table = doc.add_table(rows=6, cols=2)
    stats_table.style = 'Light Grid Accent 2'
    
    # Header
    stats_header = stats_table.rows[0].cells
    stats_header[0].text = 'Content Type'
    stats_header[1].text = 'Expected Count'
    
    # Data
    for i, (stat_type, stat_value) in enumerate(stats, 1):
        stats_row = stats_table.rows[i].cells
        stats_row[0].text = stat_type
        stats_row[1].text = stat_value
    
    # Add conclusion
    doc.add_paragraph()
    doc.add_heading('Conclusion', 1)
    doc.add_paragraph(
        'This test content demonstrates the type of information that will be available '
        'in the CloudFuze chatbot knowledge base once the SharePoint integration is fully implemented.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The SharePoint DOC360 site contains comprehensive documentation about migration capabilities, '
        'feature compatibility matrices, FAQs, and detailed feature definitions that will significantly '
        'enhance the chatbot\'s ability to answer user questions about CloudFuze migration services.'
    )
    
    # Save the document
    output_file = 'SharePoint_Test_Content.docx'
    doc.save(output_file)
    
    print(f"✅ Word document created: {output_file}")
    print(f"   Location: {os.path.abspath(output_file)}")
    
    return output_file

if __name__ == "__main__":
    import os
    output_file = create_sharepoint_test_report()
    print(f"\n📄 Document ready: {output_file}")
    print(f"   Total pages: ~5")
    print(f"   Content sections: 6")
    print(f"   Test tables: 2")
