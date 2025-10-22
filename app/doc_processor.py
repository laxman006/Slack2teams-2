import os
import pandas as pd
from typing import List, Dict
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Try to import python-docx for Word document support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available, Word document processing may be limited")

# Try to import python-docx2txt as alternative
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False
    print("docx2txt not available, using alternative methods")

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text content from a Word document (.docx file)."""
    text_content = []
    
    try:
        # Method 1: Try python-docx first (most reliable)
        if DOCX_AVAILABLE:
            try:
                doc = DocxDocument(docx_path)
                
                # Add file-level metadata for better searchability
                text_content.append(f"Word document: {os.path.basename(docx_path)}")
                text_content.append(f"Contains structured text and formatting information")
                
                # Extract text from all paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_data.append(cell_text)
                        if row_data:
                            text_content.append(" | ".join(row_data))
                
                return "\n".join(text_content)
                
            except Exception as e:
                print(f"python-docx failed for {docx_path}: {e}")
        
        # Method 2: Try docx2txt as fallback
        if DOCX2TXT_AVAILABLE:
            try:
                text = docx2txt.process(docx_path)
                if text.strip():
                    return f"Word document: {os.path.basename(docx_path)}\n\n{text}"
            except Exception as e:
                print(f"docx2txt failed for {docx_path}: {e}")
        
        # Method 3: Try to read as plain text (for some .docx files)
        try:
            with open(docx_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                if content.strip():
                    return f"Word document: {os.path.basename(docx_path)}\n\n{content}"
        except Exception as e:
            print(f"Plain text reading failed for {docx_path}: {e}")
        
        return ""
        
    except Exception as e:
        print(f"Error reading Word document {docx_path}: {e}")
        return ""

def process_doc_directory(doc_directory: str) -> List[Document]:
    """Process all Word documents in a directory and return as LangChain Documents."""
    documents = []
    
    if not os.path.exists(doc_directory):
        print(f"Word documents directory {doc_directory} does not exist")
        return documents
    
    # Supported Word document file extensions
    doc_extensions = ['.docx', '.doc']
    doc_files = []
    
    for file in os.listdir(doc_directory):
        if any(file.lower().endswith(ext) for ext in doc_extensions):
            doc_files.append(file)
    
    if not doc_files:
        print(f"No Word documents found in {doc_directory}")
        return documents
    
    print(f"Processing {len(doc_files)} Word document(s)...")
    
    for doc_file in doc_files:
        doc_path = os.path.join(doc_directory, doc_file)
        print(f"Processing: {doc_file}")
        
        try:
            # Extract text from Word document
            text = extract_text_from_docx(doc_path)
            source_type = "doc"
            
            if text.strip():
                # Create a document with metadata
                doc = Document(
                    page_content=text,
                    metadata={
                        "source": doc_file,
                        "source_type": source_type,
                        "file_path": doc_path,
                        "file_format": doc_file.split('.')[-1].lower(),
                        "content_type": "word_document",
                        "searchable_terms": " ".join(text.split()[:20])  # Add first 20 words for better searchability
                    }
                )
                documents.append(doc)
                print(f"Successfully processed {doc_file} ({len(text)} characters)")
            else:
                print(f"Warning: No text extracted from {doc_file}")
                
        except Exception as e:
            print(f"Error processing {doc_file}: {e}")
    
    return documents

def chunk_doc_documents(documents: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    """Split Word documents into smaller chunks for better retrieval."""
    if not documents:
        return documents
    
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " | ", " ", ""]  # Better separators for Word documents
    )
    
    chunked_docs = []
    for doc in documents:
        chunks = splitter.split_documents([doc])
        # Add metadata to each chunk for better searchability
        for chunk in chunks:
            chunk.metadata.update({
                "chunk_type": "word_document",
                "searchable_content": " ".join(chunk.page_content.split()[:20])  # First 20 words for search
            })
        chunked_docs.extend(chunks)
    
    print(f"Split {len(documents)} Word documents into {len(chunked_docs)} chunks")
    return chunked_docs

def get_doc_summary(doc_path: str) -> Dict:
    """Get a summary of a Word document's structure."""
    try:
        if DOCX_AVAILABLE:
            doc = DocxDocument(doc_path)
            summary = {
                "file_name": os.path.basename(doc_path),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "total_text_length": 0
            }
            
            # Count text length
            for paragraph in doc.paragraphs:
                summary["total_text_length"] += len(paragraph.text)
            
            return summary
        else:
            return {"error": "python-docx not available for detailed analysis"}
    except Exception as e:
        return {"error": str(e)}
